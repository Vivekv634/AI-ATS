"""
DOCX/DOC document text extractor.

Uses python-docx for DOCX files.
"""

import io
from pathlib import Path

from src.utils.logger import get_logger

from .base import BaseExtractor, ExtractionResult

logger = get_logger(__name__)


class DOCXExtractor(BaseExtractor):
    """Extractor for Microsoft Word documents (.docx, .doc)."""

    @property
    def supported_extensions(self) -> tuple[str, ...]:
        return (".docx", ".doc")

    def extract(self, file_path: str | Path) -> ExtractionResult:
        """Extract text from a DOCX/DOC file."""
        try:
            path = self._validate_file(file_path)

            if path.suffix.lower() == ".doc":
                return self._extract_doc(path)
            else:
                return self._extract_docx(path)

        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path}: {e}")
            return self._create_error_result(e)

    def extract_from_bytes(
        self, content: bytes, filename: str = "document.docx"
    ) -> ExtractionResult:
        """Extract text from DOCX bytes."""
        try:
            file_obj = io.BytesIO(content)
            ext = Path(filename).suffix.lower()

            if ext == ".doc":
                # .doc from bytes is tricky, save to temp file
                return ExtractionResult(
                    text="",
                    success=False,
                    error_message="DOC format from bytes not supported, use DOCX",
                )

            return self._extract_docx_from_file_obj(file_obj)

        except Exception as e:
            logger.error(f"DOCX extraction from bytes failed: {e}")
            return self._create_error_result(e)

    def _extract_docx(self, file_path: Path) -> ExtractionResult:
        """Extract text from a DOCX file."""
        try:
            from docx import Document

            doc = Document(str(file_path))
            return self._process_document(doc)

        except ImportError:
            logger.error("python-docx not installed")
            return ExtractionResult(
                text="",
                success=False,
                error_message="python-docx library not installed",
            )

    def _extract_docx_from_file_obj(self, file_obj) -> ExtractionResult:
        """Extract text from a DOCX file object."""
        try:
            from docx import Document

            doc = Document(file_obj)
            return self._process_document(doc)

        except ImportError:
            logger.error("python-docx not installed")
            return ExtractionResult(
                text="",
                success=False,
                error_message="python-docx library not installed",
            )

    def _process_document(self, doc) -> ExtractionResult:
        """Process a python-docx Document object."""
        text_parts = []
        metadata = {"extractor": "python-docx"}
        warnings = []

        # Extract core properties
        try:
            if doc.core_properties:
                props = doc.core_properties
                metadata["document_properties"] = {
                    "author": props.author,
                    "title": props.title,
                    "subject": props.subject,
                    "created": str(props.created) if props.created else None,
                    "modified": str(props.modified) if props.modified else None,
                }
        except Exception as e:
            logger.debug(f"Could not extract document properties: {e}")

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)

        # Extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_texts.append(" | ".join(row_text))

        if table_texts:
            text_parts.append("\n--- Tables ---\n")
            text_parts.extend(table_texts)

        # Count sections as "pages" (approximate)
        section_count = len(doc.sections) if doc.sections else 1

        full_text = "\n".join(text_parts)

        if not full_text.strip():
            warnings.append("Document appears to be empty or contains only images")

        return ExtractionResult(
            text=full_text,
            page_count=section_count,
            metadata=metadata,
            warnings=warnings,
        )

    def _extract_doc(self, file_path: Path) -> ExtractionResult:
        """
        Extract text from legacy .doc files.

        Note: This requires additional system dependencies or antiword.
        Falls back to basic text extraction.
        """
        warnings = ["Legacy .doc format - extraction may be incomplete"]

        # Try using textract if available
        try:
            import textract
            text = textract.process(str(file_path)).decode("utf-8")
            return ExtractionResult(
                text=text,
                page_count=1,
                metadata={"extractor": "textract"},
                warnings=warnings,
            )
        except ImportError:
            pass
        except Exception as e:
            logger.debug(f"textract failed: {e}")

        # Try basic binary extraction as last resort
        try:
            text = self._basic_doc_extract(file_path)
            if text:
                warnings.append("Used basic extraction - formatting may be lost")
                return ExtractionResult(
                    text=text,
                    page_count=1,
                    metadata={"extractor": "basic"},
                    warnings=warnings,
                )
        except Exception as e:
            logger.debug(f"Basic DOC extraction failed: {e}")

        return ExtractionResult(
            text="",
            success=False,
            error_message="Could not extract text from .doc file. Convert to .docx for better results.",
            warnings=warnings,
        )

    def _basic_doc_extract(self, file_path: Path) -> str:
        """
        Basic text extraction from .doc by reading binary.
        This is a fallback that extracts visible ASCII text.
        """
        import re

        with open(file_path, "rb") as f:
            content = f.read()

        # Try to decode and extract readable text
        try:
            # Look for text between common DOC markers
            text = content.decode("utf-8", errors="ignore")
        except Exception:
            text = content.decode("latin-1", errors="ignore")

        # Clean up binary artifacts
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", " ", text)
        text = re.sub(r"\s+", " ", text)

        # Extract words (filter out gibberish)
        words = []
        for word in text.split():
            # Keep words that look like real words
            if len(word) >= 2 and word.isalnum() or any(c.isalpha() for c in word):
                if not all(c in "0123456789" for c in word):  # Skip pure numbers
                    words.append(word)

        return " ".join(words[:5000])  # Limit output
